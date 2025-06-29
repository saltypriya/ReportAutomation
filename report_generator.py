import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_HEADER_FOOTER
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, Alignment, Font
from datetime import datetime


def set_custom_styles(doc):
    styles = doc.styles
    if 'PhotoCaption' not in styles:
        caption_style = styles.add_style('PhotoCaption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = 'Calibri'
        caption_style.font.size = Pt(9)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(120, 120, 120)


def add_header_footer_with_images(doc, photos_path):
    header_path = os.path.join(photos_path, "header.png")
    footer_path = os.path.join(photos_path, "footer.png")
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        if os.path.exists(header_path):
            run = header.paragraphs[0].add_run()
            run.add_picture(header_path, width=Inches(6.5))
            header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer = section.footer
        footer.is_linked_to_previous = False
        if os.path.exists(footer_path):
            run = footer.paragraphs[0].add_run()
            run.add_picture(footer_path, width=Inches(6.5))
            footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_excel_borders(sheet):
    thin = Side(style='thin')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    label_keywords = [
        "INSURED", "ADDRESS", "INSURER", "CLAIM #",
        "ADJUSTER", "DATE OF INSPECTION", "DATE OF LOSS", "DATE OF REPORT",
        "TYPE OF LOSS", "CAUSE OF LOSS", "SCOPE OF WORK", "CONCLUSION",
        "RECOMMENDED", "THANK YOU"
    ]

    for row_idx, row in enumerate(sheet.iter_rows(), start=1):
        for col_idx, cell in enumerate(row, start=1):
            cell.border = border
            cell.alignment = Alignment(wrap_text=True, vertical='top')
            value = str(cell.value).upper().strip() if cell.value else ""
            if row_idx == 1 and col_idx == 1:
                cell.font = Font(name='Calibri', size=14, bold=True, italic=True)
            elif any(kw in value for kw in label_keywords):
                cell.font = Font(name='Calibri', size=11, bold=True)
            else:
                cell.font = Font(name='Calibri', size=11)


def format_date(value):
    return value.strftime('%B %d, %Y').upper() if isinstance(value, datetime) else str(value).upper()


def add_image_with_caption(doc, image_path, width, height, caption_text):
    run = doc.add_paragraph().add_run()
    run.add_picture(image_path, width=width, height=height)
    caption = doc.add_paragraph(caption_text, style='PhotoCaption')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_room_images(doc, room_name, image_paths, start_count, width, height):
    doc.add_page_break()
    doc.add_paragraph(room_name).runs[0].bold = True

    image_count = start_count
    rows = (len(image_paths) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.columns[0].width = width
    table.columns[1].width = width

    for i, img in enumerate(image_paths[:4]):
        cell = table.cell(i // 2, i % 2)
        run = cell.paragraphs[0].add_run()
        run.add_picture(img, width=width, height=height)
        caption = cell.add_paragraph(f"Image {image_count}", style='PhotoCaption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_count += 1

    return image_count


def generate_perfect_reports(excel_path, output_folder, photos_path="photos"):
    wb = load_workbook(excel_path)
    sheet = wb.active
    add_excel_borders(sheet)
    wb.save(excel_path)

    os.makedirs(output_folder, exist_ok=True)

    for row in sheet.iter_rows(min_row=2, values_only=True):
        (insured_name, address, insurer, claim_number, doi, dol, dor,
         loss_type, cause_loss, indemnity, listing, contents) = row

        if not all([insured_name, address, claim_number]):
            continue

        doi, dol, dor = format_date(doi), format_date(dol), format_date(dor)

        doc = Document()
        set_custom_styles(doc)
        add_header_footer_with_images(doc, photos_path)

        image_counter = 1
        width = Inches(3.25)
        height = Inches(2.25)

        doc.add_paragraph("FIRST INSPECTION REPORT").alignment = WD_ALIGN_PARAGRAPH.CENTER

        header_lines = [
            f"INSURED/POLICYHOLDER: {insured_name}",
            f"ADDRESS: {address}",
            f"INSURER: {insurer}",
            f"CLAIM #: {claim_number}",
            f"ADJUSTER/ CLAIM REP: TOP GUN",
            f"DATE OF INSPECTION: {doi}",
            f"DATE OF LOSS: {dol}",
            f"DATE OF REPORT: {dor}",
            f"TYPE OF LOSS: {loss_type.upper()}"
        ]
        for line in header_lines:
            doc.add_paragraph(line)

        front_img = os.path.join(photos_path, "front_house.jpg")
        if os.path.exists(front_img):
            add_image_with_caption(doc, front_img, width, height, f"Image {image_counter}")
            image_counter += 1

        doc.add_paragraph("CAUSE OF LOSS:").runs[0].bold = True
        doc.add_paragraph(str(cause_loss).strip())

        doc.add_paragraph("SCOPE OF WORK:").runs[0].bold = True
        doc.add_paragraph("The following is a brief outline of the work to be completed on the contents portion of this claim.")
        for i, item in enumerate([
            "Assess, pack and move out all salvageable contents.",
            "Inventory all the affected contents.",
            "Inspect all affected electronics.",
            "Restore salvageable contents.",
            "Dispose of non-salvageable contents."
        ], start=1):
            doc.add_paragraph(f"{i}. {item}")

        doc.add_paragraph("RECOMMENDED RESERVES FOR TRINITY’S INVOLVEMENT:").runs[0].bold = True
        doc.add_paragraph("The estimated cost for Trinity’s involvement is as follows:")
        doc.add_paragraph(f"• Indemnity Work: Should not exceed ${indemnity:,.2f} plus HST")
        doc.add_paragraph("Our actual cost will be adjusted once the exact scope of approved work is known. "
                          "The recommended estimate is only based on visual inspection for reserves setting purposes.")
        doc.add_paragraph(f"• Trinity Listing & Pricing Expense Reserve: Should not exceed ${listing:,.2f} plus HST")

        doc.add_paragraph("RECOMMENDED RESERVES FOR THE TOTAL CONTENTS LOSS:").runs[0].bold = True
        doc.add_paragraph(f"Based on a visual inspection of the extent of non-salvageable items on the main floor, "
                          f"we believe that the total replacement cost for the non-salvageable items should not exceed "
                          f"${contents:,.2f} plus HST.")

        doc.add_paragraph("CONCLUSION:").runs[0].bold = True
        doc.add_paragraph("Once our scope of work is approved, we can attend and begin the pack out process.")

        doc.add_paragraph("\nThank You,\n\nMo Waez\nTrinity Contents Management\nmo@trinitycontents.com\n(647) 613-2246")

        room_config = [
            ("KITCHEN & DINING AREA", "kitchen"),
            ("LIVING ROOM", "living"),
            ("BEDROOM 1", "bedroom1"),
            ("BEDROOM 2", "bedroom2"),
            ("STORAGE ROOM", "storage")
        ]
        for room_name, folder in room_config:
            folder_path = os.path.join(photos_path, folder)
            if os.path.exists(folder_path):
                images = sorted([
                    os.path.join(folder_path, f)
                    for f in os.listdir(folder_path)
                    if f.lower().endswith(('.jpg', '.jpeg'))
                ])
                if images:
                    image_counter = add_room_images(doc, room_name, images, image_counter, width, height)

        last = insured_name.split()[-1]
        street = address.split(',')[0].strip()
        filename = f"FIRST INSPECTION REPORT - CLAIM# {claim_number} - {last} - {street}.docx"
        safe_filename = "".join(c for c in filename if c not in '\\/:*?"<>|')
        doc.save(os.path.join(output_folder, safe_filename))
